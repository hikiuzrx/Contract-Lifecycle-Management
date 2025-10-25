import os
from pathlib import Path
from typing import Optional
from docx import Document
from pdfminer.high_level import extract_text as pdfminer_extract_text
import pytesseract
from PIL import Image
import pypdfium2 as pdfium
from app.config import get_config, Config 


class DocumentExtractor:
    """
    A worker class responsible for extracting text from documents.
    It supports DOCX and PDF (with digital text preference and OCR fallback).
    """
    def __init__(self, config: Config):
        self.config = config
        
        # 1. Configure Tesseract Path
        if self.config.ocr.tesseract_path:
            pytesseract.pytesseract.tesseract_cmd = self.config.ocr.tesseract_path

    def extract(self, file_path: Path) -> str:
        """
        Main method to extract text based on file extension.
        
        Args:
            file_path: The Path object pointing to the document.

        Returns:
            The extracted text content or an error message.
        """
        file_path = Path(file_path)
        
        if file_path.suffix.lower() == '.pdf':
            return self._extract_pdf(file_path)
        elif file_path.suffix.lower() == '.docx':
            return self._extract_docx(file_path)
        else:
            return f"Error: Unsupported file format {file_path.suffix}."

    def _extract_docx(self, file_path: Path) -> str:
        print(f"Processing digital DOCX file: {file_path.name}")
        try:
            document = Document(file_path)
            full_text = []
            
            for paragraph in document.paragraphs:
                full_text.append(paragraph.text)
            
            for table in document.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    full_text.append(f"\n---TABLE ROW---\n{row_text}")

            return "\n\n".join(full_text)
        except Exception as e:
            return f"Error extracting DOCX text: {e}"

    def _extract_pdf(self, file_path: Path) -> str:
        """Extracts text from a PDF, falling back to OCR if digital text is insufficient."""
        
        ocr_config = self.config.ocr
        
        # 1. Try Digital Extraction
        try:
            digital_text = pdfminer_extract_text(file_path)
            
            # Check if digital text length is sufficient (i.e., not a scanned document)
            if len(digital_text.strip()) >= ocr_config.min_text_length:
                print(f"PDF is digital (length {len(digital_text.strip())}), using fast extraction.")
                return digital_text
            
            print(f"Digital text length ({len(digital_text.strip())}) is too short. Falling back to OCR.")
        except Exception as e:
            # Catch exceptions from pdfminer (e.g., malformed PDF)
            print(f"Digital extraction failed ({e}). Falling back to OCR.")

        # 2. OCR Extraction
        return self._ocr_pdf(file_path)

    def _ocr_pdf(self, file_path: Path) -> str:
        """Performs Tesseract OCR on a PDF file using configuration settings."""
        
        ocr_config = self.config.ocr
        
        # Tesseract arguments, using psm_mode 6 and eng+ara language
        tesseract_args = (
            f'--oem 3 --psm {ocr_config.psm_mode} -l {ocr_config.language}'
        )
        full_ocr_text = []
        
        try:
            # Load PDF document
            pdf_document = pdfium.PdfDocument(file_path)
            
            # Calculate scale factor for DPI (e.g., 300 dpi / 72 base dpi)
            scale = ocr_config.dpi / 72.0 
            
            for page_index in range(len(pdf_document)):
                print(f"Processing Page {page_index + 1} with OCR...")
                page = pdf_document.get_page(page_index)
                
                # Render page to a PIL Image object
                bitmap = page.render(scale=scale)
                pil_image = bitmap.to_pil()
                
                # Run Tesseract OCR with the configured parameters
                page_text = pytesseract.image_to_string(pil_image, config=tesseract_args)
                full_ocr_text.append(page_text)
                
            return "\n\n---PAGE BREAK---\n\n".join(full_ocr_text)

        except pytesseract.TesseractNotFoundError:
            # Handle Tesseract not found error, which is critical for OCR
            return f"Error: Tesseract not found. Check TESSERACT_PATH in config and ensure Tesseract is installed."
        except Exception as e:
            return f"Error during PDF OCR extraction: {e}"


# --- Worker Function ---

def document_extraction_worker(file_path: str) -> str:
    """
    The main worker entry point.
    
    In a real system, this would be called by a task queue (like Celery) 
    with a path to a file in a storage bucket.
    """
    file_path_obj = Path(file_path)
    print(f"Starting extraction worker for file: {file_path_obj.name}")
    
    try:
        # Get the global, validated configuration
        config = get_config() 
        
        # Initialize the extractor with the config
        extractor = DocumentExtractor(config)
        extracted_text = extractor.extract(file_path_obj)
        
        return extracted_text
    
    except Exception as e:
        print(f"CRITICAL WORKER FAILURE: {e}")
        return f"Extraction failed: {e}"

